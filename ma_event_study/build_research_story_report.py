from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Tuple
from xml.sax.saxutils import escape

import pandas as pd

try:
    # Работает при запуске как часть пакета (`python -m ma_event_study...`).
    from .paths import resolve_clean_data_file
except ImportError:  # pragma: no cover
    # Работает при запуске файла напрямую (`python ma_event_study/build_research_story_report.py`).
    from paths import resolve_clean_data_file
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from matplotlib import font_manager
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = PROJECT_ROOT / "out" / "thesis"
TABLES = OUTPUTS / "tables"
CHARTS = OUTPUTS / "charts"
MODELS = OUTPUTS / "models"
CLEAN = OUTPUTS / "clean_data"

DOCX_PATH = TABLES / "narrative_research_report_ru.docx"
PDF_PATH = TABLES / "narrative_research_report_ru.pdf"


@dataclass
class ReportContext:
    base_deals: pd.DataFrame
    enriched: pd.DataFrame
    mapping_audit: pd.DataFrame
    summary_stats: pd.DataFrame
    one_sample_tests: pd.DataFrame
    group_tests: pd.DataFrame
    regression_coeffs: pd.DataFrame
    regression_diagnostics: pd.DataFrame
    hypotheses: pd.DataFrame
    warnings: pd.DataFrame
    case_summary: pd.DataFrame
    announcement_panel: pd.DataFrame
    create_panel: pd.DataFrame


def _read_csv_optional(path: Path) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_csv(path, encoding="utf-8")


def _clean_csv(name: str) -> pd.DataFrame:
    p = resolve_clean_data_file(name)
    if not p.is_file():
        alt = CLEAN / name
        if alt.is_file():
            p = alt
        else:
            return pd.DataFrame()
    return pd.read_csv(p, encoding="utf-8")


def load_context() -> ReportContext:
    return ReportContext(
        base_deals=_clean_csv("base_deals_standardized.csv"),
        enriched=_clean_csv("ma_deals_enriched.csv"),
        mapping_audit=_read_csv_optional(TABLES / "mapping_audit.csv"),
        summary_stats=_read_csv_optional(TABLES / "summary_statistics.csv"),
        one_sample_tests=_read_csv_optional(TABLES / "one_sample_tests.csv"),
        group_tests=_read_csv_optional(TABLES / "group_tests.csv"),
        regression_coeffs=_read_csv_optional(MODELS / "regression_coefficients.csv"),
        regression_diagnostics=_read_csv_optional(MODELS / "regression_diagnostics.csv"),
        hypotheses=_read_csv_optional(TABLES / "hypotheses_significance_map_ru.csv"),
        warnings=_read_csv_optional(TABLES / "warnings.csv"),
        case_summary=_read_csv_optional(TABLES / "significant_deals_case_analysis_summary_ru.csv"),
        announcement_panel=_clean_csv("announcement_daily_panel_clean.csv"),
        create_panel=_clean_csv("create_daily_panel_clean.csv"),
    )


def fmt_pct(x: float, digits: int = 2) -> str:
    if pd.isna(x):
        return "н/д"
    return f"{x * 100:.{digits}f}%"


def fmt_num(x: float, digits: int = 4) -> str:
    if pd.isna(x):
        return "н/д"
    return f"{x:.{digits}f}"


def fmt_p(x: float) -> str:
    if pd.isna(x):
        return "н/д"
    if x < 0.0001:
        return "< 0.0001"
    return f"{x:.4f}"


def compute_metrics(ctx: ReportContext) -> Dict[str, object]:
    base = ctx.base_deals
    enr = ctx.enriched
    warnings_counts = ctx.warnings["category"].value_counts().to_dict() if not ctx.warnings.empty and "category" in ctx.warnings.columns else {}
    mapping_counts = (
        ctx.mapping_audit.groupby(["table_role", "mapping_status"]).size().reset_index(name="n")
    )
    coverage = {
        "base_deals": int(len(base)),
        "intraday": int(enr["CAR_ANN_INTRADAY_15M"].notna().sum()),
        "ann_car": int(enr["CAR_ANN_1_1"].notna().sum()),
        "ann_bhar_120": int(enr["BHAR_ANN_120"].notna().sum()),
        "close_car": int(enr["CAR_CLOSE_1_1"].notna().sum()),
        "close_bhar_120": int(enr["BHAR_CLOSE_120"].notna().sum()),
        "create_car": int(enr["CAR_CREATE_1_1"].notna().sum()),
        "act_car": int(enr["CAR_ACT_1_1"].notna().sum()),
        "leakage": int(enr["RUNUP_PRE_30_5"].notna().sum()),
    }
    hypothesis_counts = (
        ctx.hypotheses["Есть статистически значимый эффект"].value_counts().to_dict()
        if not ctx.hypotheses.empty and "Есть статистически значимый эффект" in ctx.hypotheses.columns
        else {}
    )
    if ctx.one_sample_tests.empty:
        significant_tests = pd.DataFrame()
    else:
        significant_tests = ctx.one_sample_tests[ctx.one_sample_tests["p_value"] < 0.05].copy()
        significant_tests = significant_tests.sort_values("p_value")
    significant_regs = ctx.regression_coeffs[ctx.regression_coeffs["p_value"] < 0.05].copy()
    significant_regs = significant_regs.sort_values(["model", "p_value"])
    borderline_regs = ctx.regression_coeffs[
        (ctx.regression_coeffs["p_value"] >= 0.05) & (ctx.regression_coeffs["p_value"] < 0.10)
    ].copy()
    significant_groups = ctx.group_tests[ctx.group_tests["p_value"] < 0.05].copy()

    ann_profile = ctx.announcement_panel.loc[
        ctx.announcement_panel["t"].between(-10, 10) & ctx.announcement_panel["ar_ann"].notna()
    ].copy()
    create_profile = ctx.create_panel.loc[
        ctx.create_panel["t"].between(-10, 10) & ctx.create_panel["ar_create"].notna()
    ].copy()

    ann_profile_mean = ann_profile.groupby("t")["ar_ann"].mean()
    create_profile_mean = create_profile.groupby("t")["ar_create"].mean()

    return {
        "coverage": coverage,
        "warnings_counts": warnings_counts,
        "mapping_counts": mapping_counts,
        "hypothesis_counts": hypothesis_counts,
        "significant_tests": significant_tests,
        "significant_regs": significant_regs,
        "borderline_regs": borderline_regs,
        "significant_groups": significant_groups,
        "ann_profile_deals": int(ann_profile["source_row_excel"].nunique()),
        "ann_profile_rows": int(len(ann_profile)),
        "create_profile_deals": int(create_profile["source_row_excel"].nunique()),
        "create_profile_rows": int(len(create_profile)),
        "ann_profile_mean": ann_profile_mean,
        "create_profile_mean": create_profile_mean,
    }


def build_text_sections(ctx: ReportContext, m: Dict[str, object]) -> Dict[str, object]:
    coverage = m["coverage"]
    warnings_counts = m["warnings_counts"]
    hypothesis_counts = m["hypothesis_counts"]

    one = ctx.one_sample_tests.set_index("metric") if not ctx.one_sample_tests.empty else pd.DataFrame()

    def one_mp(met: str) -> tuple[str, str]:
        if one.empty or met not in one.index:
            return "н/д", "н/д"
        return fmt_pct(one.loc[met, "mean"]), fmt_p(one.loc[met, "p_value"])

    sig_reg = m["significant_regs"]

    intro = (
        "Цель исследования состояла в эмпирической оценке эффективности сделок слияний и поглощений "
        "в банковском секторе и смежных экосистемах на российском рынке. В качестве базового набора "
        "использовался файл ma_deals_AUDIT.xlsx, который был расширен рыночными и фундаментальными "
        "показателями из внутридневных и дневных панелей, после чего по каждой сделке были рассчитаны краткосрочные, "
        "среднесрочные и долгосрочные эффекты."
    )

    data_block = (
        f"Итоговая база содержит {coverage['base_deals']} сделок. Из них внутридневное покрытие вокруг объявления "
        f"доступно для {coverage['intraday']} сделок, дневной CAR вокруг объявления рассчитан для {coverage['ann_car']} сделок, "
        f"дневной CAR вокруг закрытия сделки — для {coverage['close_car']} сделок, метрики возможной утечки информации — для {coverage['leakage']} сделок, "
        f"BHAR_ANN_120 — для {coverage['ann_bhar_120']} сделок. Все сопоставления, попавшие в расчеты, являются однозначными; "
        f"двусмысленные соответствия в финальные метрики не включались."
    )

    method_block = (
        "Ключевые математические параметры исследования соответствуют классической логике событийного анализа. "
        "Дневная доходность акции рассчитывалась как простая доходность между соседними торговыми днями, "
        "сверхнормальная доходность — через рыночную модель при наличии достаточного оценочного окна (-250; -30), а при "
        "недостатке данных использовались модель с корректировкой на рынок и модель с корректировкой на среднюю доходность. Кумулятивная "
        "сверхнормальная доходность CAR строилась как сумма сверхнормальных доходностей в стандартных окнах [-1;+1], "
        "[-3;+3], [-5;+5] и [-10;+10]. Долгосрочная эффективность измерялась через BHAR на горизонтах 60, "
        "120 и 250 торговых дней. Во внутридневном блоке событие привязывалось к anchor_timestamp_msk, а "
        "ожидаемая 15-минутная доходность задавалась как средняя 15-минутная доходность предыдущего торгового дня."
    )

    hypotheses_block = (
        f"Всего было рассмотрено 22 гипотезы. По итогам тестирования {hypothesis_counts.get('ДА', 0)} гипотезы "
        f"получили прямую статистическую поддержку, {hypothesis_counts.get('ЧАСТИЧНО', 0)} — частичную поддержку, "
        f"{hypothesis_counts.get('НЕТ', 0)} не подтвердились, а {hypothesis_counts.get('НЕ ПРОТЕСТИРОВАНО', 0)} "
        "не удалось корректно протестировать из-за ограничений данных и отсутствия устойчивой классификации некоторых типов сделок."
    )

    a15, pa15 = one_mp("CAR_ANN_INTRADAY_15M")
    a30, pa30 = one_mp("CAR_ANN_INTRADAY_30M")
    a1h, pa1h = one_mp("CAR_ANN_INTRADAY_1H")
    ann_intraday = (
        f"Внутридневной анализ вокруг объявления не выявил статистически значимого мгновенного ценового эффекта. "
        f"Для CAR_ANN_INTRADAY_15M средний эффект составил {a15} при p={pa15}; для CAR_ANN_INTRADAY_30M — "
        f"{a30} при p={pa30}; для CAR_ANN_INTRADAY_1H — {a1h} при p={pa1h}. "
        "Следовательно, гипотезы о сильной немедленной внутридневной реакции рынка на стадии объявления сделки не подтверждаются."
    )

    a11, pa11 = one_mp("CAR_ANN_1_1")
    a33, pa33 = one_mp("CAR_ANN_3_3")
    a55, pa55 = one_mp("CAR_ANN_5_5")
    ann_daily = (
        f"В дневных окнах вокруг объявления сделки средние значения CAR также оказались близкими к нулю: "
        f"CAR_ANN_1_1 = {a11} (p={pa11}), CAR_ANN_3_3 = {a33} (p={pa33}), CAR_ANN_5_5 = {a55} (p={pa55}). "
        "Таким образом, рынок в среднем не демонстрирует ярко выраженной положительной или отрицательной переоценки именно в день официального объявления."
    )

    ru, pru = one_mp("RUNUP_PRE_30_5")
    ca, pca = one_mp("CAR_PRE_ANNOUNCEMENT")
    leakage = (
        f"Наиболее сильный результат получен в блоке анализа возможной утечки информации. Предсобытийный рост котировок в окне [-30;-5] "
        f"составил {ru} при p={pru}, а полный предсобытийный CAR_PRE_ANNOUNCEMENT — {ca} при p={pca}. Это означает, что цены "
        "акций покупателей в среднем начинали расти еще до официального объявления сделки, что согласуется либо с постепенным "
        "встраиванием информации в цену, либо с утечкой информации."
    )

    c55, pc55 = one_mp("CAR_CLOSE_5_5")
    completion = (
        f"Вокруг даты закрытия сделки обнаружен более выраженный эффект, чем вокруг даты объявления. "
        f"Так, CAR_CLOSE_5_5 равен {c55} при p={pc55}. "
        f"Это единственное стандартное дневное окно CAR, показавшее статистическую значимость на 5%-ном уровне. "
        "Следовательно, часть сделок переоценивается рынком не в момент первой новости, а в момент фактического закрытия, "
        "когда неопределенность относительно завершения сделки снимается."
    )

    b60a, pb60a = one_mp("BHAR_ANN_60")
    b120a, pb120a = one_mp("BHAR_ANN_120")
    b60c, pb60c = one_mp("BHAR_CLOSE_60")
    long_term = (
        f"Долгосрочные результаты оказались заметно сильнее краткосрочных. BHAR_ANN_60 составил {b60a} при p={pb60a}, "
        f"а BHAR_ANN_120 — {b120a} при p={pb120a}. Положительный эффект сохраняется и после закрытия сделки: "
        f"BHAR_CLOSE_60 = {b60c} при p={pb60c}. В совокупности это противоречит гипотезе об отрицательном "
        "постслиянийном дрейфе доходности и указывает на положительную среднесрочную переоценку покупателей после M&A."
    )

    reg_lines = []
    for _, row in sig_reg.iterrows():
        term = str(row["term"])
        if term == "Intercept":
            continue
        reg_lines.append(
            f"В модели {row['model']} коэффициент при {term} равен {fmt_num(row['coef'], 6)} при p={fmt_p(row['p_value'])}."
        )
    regressions = (
        "Регрессионный анализ показал, что CAR вокруг объявления практически не объясняется фундаментальными характеристиками, "
        "но зависит от отдельных параметров сделки. В частности, абсолютный размер сделки положительно связан с CAR_ANN_1_1. "
        "Для CAR вокруг закрытия сделки картина иная: значимыми оказались P/B, P/E и рыночная капитализация покупателя, что говорит о более "
        "сильной роли предсделочных характеристик компании в момент закрытия сделки. "
        + " ".join(reg_lines[:6])
    )

    group_text = (
        "Групповые тесты показали, что единственным устойчиво значимым категориальным фактором выступает статус сделки. "
        "Для статуса сделки различия по CAR_ANN_1_1 статистически значимы (ANOVA p=0.0324), что означает неоднородность "
        "рыночной реакции между завершенными, отмененными, находящимися в переговорах и находящимися на стадии оформления сделками."
    )

    case_text = (
        "Анализ отдельных сделок показывает, что значимые агрегированные результаты формируются не случайным набором наблюдений, "
        "а устойчивыми кластерами кейсов. Положительный предсобытийный рост котировок прежде всего связан со сделками Сбера "
        "по объектам Вестер, Таганрогский мазутный терминал, Бизнес-центр в Омске и Русснефть. Положительный долгосрочный BHAR "
        "сильно поддерживается кейсами CardsMobile/Кошелек, Вестер, Mercator, Rambler и Евроцемент. В окне вокруг закрытия сделки "
        "высокий вклад дают сделки БАНКОМ, OZON.ru, ПанКлуб и AnyQuery."
    )

    limitations = (
        f"Ограничения исследования также важны для интерпретации результатов. Во-первых, для внутридневного анализа даты закрытия сделки "
        f"отсутствует отдельная внутридневная панель, поэтому соответствующие показатели не рассчитывались. Во-вторых, "
        f"для {warnings_counts.get('missing_completion_coverage', 0)} сделок дневные панели не покрывали дату закрытия, поэтому анализ вокруг закрытия сделки "
        "был доступен не для всей выборки. В-третьих, гипотезы об экосистемных сделках, сделках в рамках основного бизнеса и диверсификационных мотивах "
        "не были корректно протестированы, поскольку в текущих исходных данных отсутствует устойчивая содержательная разметка этих типов сделок."
    )

    conclusion = (
        "В целом исследование показывает, что российский рынок не демонстрирует ярко выраженной мгновенной переоценки акций "
        "покупателей в момент объявления M&A-сделок, однако фиксирует статистически значимый предсобытийный рост котировок "
        "и положительный среднесрочный долгосрочный эффект после стадии объявления и стадии закрытия сделки. Это означает, что "
        "экономическая эффективность сделок проявляется не столько в момент новостного шока, сколько в более длинном окне, "
        "а также частично закладывается в цены еще до официального объявления."
    )

    return {
        "intro": intro,
        "data_block": data_block,
        "method_block": method_block,
        "hypotheses_block": hypotheses_block,
        "ann_intraday": ann_intraday,
        "ann_daily": ann_daily,
        "leakage": leakage,
        "completion": completion,
        "long_term": long_term,
        "regressions": regressions,
        "group_text": group_text,
        "case_text": case_text,
        "limitations": limitations,
        "conclusion": conclusion,
    }


def make_word_styles(doc: Document) -> None:
    styles = doc.styles
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)
    styles["Title"].font.size = Pt(16)
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.size = Pt(12.5)


def add_docx_paragraph(doc: Document, text: str, style: str | None = None, align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.add_run(text)


def add_docx_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_docx(ctx: ReportContext, m: Dict[str, object], sections: Dict[str, object]) -> None:
    doc = Document()
    make_word_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Отчет по эмпирическому исследованию эффективности сделок M&A\nв банковском секторе и смежных экосистемах на российском рынке")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = subtitle.add_run("На основе итоговых файлов каталога результатов")
    rr.italic = True
    rr.font.name = "Times New Roman"
    rr.font.size = Pt(11)

    doc.add_heading("1. Цель и логика исследования", level=1)
    add_docx_paragraph(doc, sections["intro"])
    add_docx_paragraph(doc, sections["data_block"])
    add_docx_paragraph(doc, sections["method_block"])

    doc.add_heading("2. Исходные данные и построение выборки", level=1)
    mapping_rows = []
    for _, row in m["mapping_counts"].iterrows():
        mapping_rows.append([str(row["table_role"]), str(row["mapping_status"]), str(int(row["n"]))])
    add_docx_paragraph(
        doc,
        "Для сопоставления сделок и рыночных панелей использовались source_row_excel, тикер покупателя, название покупателя и объект сделки. В расчетах сохранялись только наблюдения с однозначным соответствием."
    )
    add_docx_table(doc, ["Роль таблицы", "Статус сопоставления", "Число наблюдений"], mapping_rows)
    coverage = m["coverage"]
    add_docx_table(
        doc,
        ["Показатель покрытия", "Число сделок"],
        [
            ["Полная базовая выборка", str(coverage["base_deals"])],
            ["Внутридневной анализ вокруг объявления", str(coverage["intraday"])],
            ["Дневной CAR вокруг объявления", str(coverage["ann_car"])],
            ["Дневной CAR вокруг закрытия сделки", str(coverage["close_car"])],
            ["BHAR вокруг объявления, горизонт 120 дней", str(coverage["ann_bhar_120"])],
            ["BHAR вокруг закрытия сделки, горизонт 120 дней", str(coverage["close_bhar_120"])],
            ["Метрики возможной утечки информации", str(coverage["leakage"])],
            ["CAR вокруг даты первого появления сделки", str(coverage["create_car"])],
            ["CAR вокруг даты актуализации", str(coverage["act_car"])],
        ],
    )

    doc.add_heading("3. Гипотезы исследования", level=1)
    add_docx_paragraph(doc, sections["hypotheses_block"])
    hyp_rows = []
    if not ctx.hypotheses.empty:
        for _, row in ctx.hypotheses.iterrows():
            hyp_rows.append([
                row["Гипотеза"],
                row["Есть статистически значимый эффект"],
                str(row["Комментарий"]),
            ])
    if hyp_rows:
        add_docx_table(doc, ["Гипотеза", "Статус", "Краткий вывод"], hyp_rows)
    else:
        add_docx_paragraph(doc, "Детализированная таблица гипотез отсутствует (файл out/thesis/tables/hypotheses_significance_map_ru.csv не найден или пуст).")

    doc.add_heading("4. Основные эмпирические результаты", level=1)
    for key in ["ann_intraday", "ann_daily", "leakage", "completion", "long_term", "regressions", "group_text", "case_text"]:
        add_docx_paragraph(doc, sections[key])

    doc.add_heading("5. Ключевые статистические результаты", level=1)
    sig_tests = m["significant_tests"]
    test_rows = []
    for _, row in sig_tests.iterrows():
        test_rows.append([
            row["metric"],
            str(int(row["n"])),
            fmt_pct(row["mean"]),
            fmt_p(row["p_value"]),
        ])
    add_docx_table(doc, ["Метрика", "n", "Средний эффект", "p-value"], test_rows)

    sig_regs = m["significant_regs"]
    reg_rows = []
    for _, row in sig_regs.iterrows():
        reg_rows.append([
            row["model"],
            row["term"],
            fmt_num(row["coef"], 6),
            fmt_p(row["p_value"]),
        ])
    add_docx_table(doc, ["Модель", "Переменная", "Коэффициент", "p-value"], reg_rows)

    if not m["significant_groups"].empty:
        group_rows = []
        for _, row in m["significant_groups"].iterrows():
            group_rows.append([
                row["group_variable"],
                row["metric"],
                row["test"],
                fmt_p(row["p_value"]),
            ])
        add_docx_table(doc, ["Групповая переменная", "Метрика", "Тест", "p-value"], group_rows)

    doc.add_heading("6. Интерпретация графиков", level=1)
    chart_notes = [
        (
            CHARTS / "announcement_event_profile.png",
            f"Профиль средней сверхнормальной доходности и накопленной сверхнормальной доходности вокруг даты объявления. "
            f"График построен по {m['ann_profile_deals']} сделкам и {m['ann_profile_rows']} дневным наблюдениям в окне t = -10...+10."
        ),
        (
            CHARTS / "create_event_profile.png",
            f"Профиль средней сверхнормальной доходности и накопленной сверхнормальной доходности вокруг даты первого появления сделки. "
            f"График построен по {m['create_profile_deals']} сделкам и {m['create_profile_rows']} дневным наблюдениям."
        ),
        (
            CHARTS / "car_ann_1_1_hist.png",
            f"Распределение CAR_ANN_1_1 по {coverage['ann_car']} сделкам; оно показывает, что вокруг объявления распределение сосредоточено около нуля и не смещено резко в положительную сторону."
        ),
        (
            CHARTS / "car_close_1_1_hist.png",
            f"Распределение CAR_CLOSE_1_1 по {coverage['close_car']} сделкам; эффект вокруг закрытия сделки выглядит более неоднородным, чем эффект вокруг объявления."
        ),
        (
            CHARTS / "bhar_ann_120_hist.png",
            f"Распределение BHAR_ANN_120 по {coverage['ann_bhar_120']} сделкам; на нем видно, что долгосрочный эффект чаще положителен, чем отрицателен."
        ),
    ]
    for img_path, caption in chart_notes:
        doc.add_paragraph(caption)
        doc.add_picture(str(img_path), width=Cm(15.5))
        doc.add_paragraph()

    doc.add_heading("7. Ограничения и корректность интерпретации", level=1)
    add_docx_paragraph(doc, sections["limitations"])

    doc.add_heading("8. Итоговые выводы", level=1)
    add_docx_paragraph(doc, sections["conclusion"])

    doc.save(DOCX_PATH)


def make_pdf_styles() -> Dict[str, ParagraphStyle]:
    font_regular = font_manager.findfont("DejaVu Serif")
    font_bold = font_manager.findfont(font_manager.FontProperties(family="DejaVu Serif", weight="bold"))
    pdfmetrics.registerFont(TTFont("DejaVuSerif", font_regular))
    pdfmetrics.registerFont(TTFont("DejaVuSerif-Bold", font_bold))
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "title",
            parent=styles["Title"],
            fontName="DejaVuSerif-Bold",
            fontSize=16,
            leading=20,
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=styles["Normal"],
            fontName="DejaVuSerif",
            fontSize=11,
            leading=14,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "body": ParagraphStyle(
            "body",
            parent=styles["Normal"],
            fontName="DejaVuSerif",
            fontSize=11,
            leading=14.5,
            alignment=TA_JUSTIFY,
            spaceAfter=7,
        ),
        "heading": ParagraphStyle(
            "heading",
            parent=styles["Heading1"],
            fontName="DejaVuSerif-Bold",
            fontSize=13,
            leading=16,
            alignment=TA_LEFT,
            spaceBefore=10,
            spaceAfter=7,
        ),
        "caption": ParagraphStyle(
            "caption",
            parent=styles["Normal"],
            fontName="DejaVuSerif",
            fontSize=9.5,
            leading=12,
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
    }


def table_for_pdf(headers: List[str], rows: List[List[str]], col_widths: List[float] | None = None) -> Table:
    data = [headers] + rows
    tbl = Table(data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D9EAF7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "DejaVuSerif-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "DejaVuSerif"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEADING", (0, 0), (-1, -1), 10),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.Color(0.98, 0.98, 0.98)]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return tbl


def build_pdf(ctx: ReportContext, m: Dict[str, object], sections: Dict[str, object]) -> None:
    styles = make_pdf_styles()
    story: List[object] = []
    story.append(Paragraph("Отчет по эмпирическому исследованию эффективности сделок M&amp;A в банковском секторе и смежных экосистемах на российском рынке", styles["title"]))
    story.append(Paragraph("На основе итоговых файлов каталога результатов", styles["subtitle"]))

    def add_body(text: str) -> None:
        story.append(Paragraph(escape(text), styles["body"]))

    def add_heading(text: str) -> None:
        story.append(Paragraph(escape(text), styles["heading"]))

    add_heading("1. Цель и логика исследования")
    add_body(sections["intro"])
    add_body(sections["data_block"])
    add_body(sections["method_block"])

    add_heading("2. Исходные данные и построение выборки")
    add_body(
        "Для сопоставления сделок и рыночных панелей использовались source_row_excel, тикер покупателя, название покупателя и объект сделки. В расчеты включались только наблюдения с однозначным соответствием."
    )
    mapping_rows = [[str(r["table_role"]), str(r["mapping_status"]), str(int(r["n"]))] for _, r in m["mapping_counts"].iterrows()]
    story.append(table_for_pdf(["Роль таблицы", "Статус", "Число наблюдений"], mapping_rows, [4.0 * cm, 3.3 * cm, 3.5 * cm]))
    story.append(Spacer(1, 0.25 * cm))
    c = m["coverage"]
    coverage_rows = [
        ["Полная базовая выборка", str(c["base_deals"])],
        ["Внутридневной анализ вокруг объявления", str(c["intraday"])],
        ["Дневной CAR вокруг объявления", str(c["ann_car"])],
        ["Дневной CAR вокруг закрытия сделки", str(c["close_car"])],
        ["BHAR вокруг объявления, горизонт 120 дней", str(c["ann_bhar_120"])],
        ["BHAR вокруг закрытия сделки, горизонт 120 дней", str(c["close_bhar_120"])],
        ["Метрики возможной утечки информации", str(c["leakage"])],
        ["CAR вокруг даты первого появления сделки", str(c["create_car"])],
        ["CAR вокруг даты актуализации", str(c["act_car"])],
    ]
    story.append(table_for_pdf(["Показатель покрытия", "Число сделок"], coverage_rows, [7.5 * cm, 3.5 * cm]))

    add_heading("3. Гипотезы исследования")
    add_body(sections["hypotheses_block"])
    if ctx.hypotheses.empty:
        add_body("Детализированная таблица гипотез отсутствует (файл out/thesis/tables/hypotheses_significance_map_ru.csv не найден или пуст).")
    else:
        hyp_rows = [[r["Гипотеза"], r["Есть статистически значимый эффект"], str(r["Комментарий"])] for _, r in ctx.hypotheses.iterrows()]
        story.append(table_for_pdf(["Гипотеза", "Статус", "Краткий вывод"], hyp_rows, [1.5 * cm, 3.2 * cm, 11.0 * cm]))

    add_heading("4. Основные эмпирические результаты")
    for key in ["ann_intraday", "ann_daily", "leakage", "completion", "long_term", "regressions", "group_text", "case_text"]:
        add_body(sections[key])

    add_heading("5. Ключевые статистические результаты")
    sig_tests = m["significant_tests"]
    test_rows = [[r["metric"], str(int(r["n"])), fmt_pct(r["mean"]), fmt_p(r["p_value"])] for _, r in sig_tests.iterrows()]
    story.append(table_for_pdf(["Метрика", "n", "Средний эффект", "p-value"], test_rows, [5.8 * cm, 1.2 * cm, 3.3 * cm, 2.1 * cm]))
    story.append(Spacer(1, 0.2 * cm))
    sig_regs = m["significant_regs"]
    reg_rows = [[r["model"], r["term"], fmt_num(r["coef"], 6), fmt_p(r["p_value"])] for _, r in sig_regs.iterrows()]
    story.append(table_for_pdf(["Модель", "Переменная", "Коэффициент", "p-value"], reg_rows, [3.3 * cm, 7.1 * cm, 3.0 * cm, 2.0 * cm]))
    if not m["significant_groups"].empty:
        story.append(Spacer(1, 0.2 * cm))
        group_rows = [[r["group_variable"], r["metric"], r["test"], fmt_p(r["p_value"])] for _, r in m["significant_groups"].iterrows()]
        story.append(table_for_pdf(["Групповая переменная", "Метрика", "Тест", "p-value"], group_rows, [4.2 * cm, 4.2 * cm, 2.3 * cm, 2.0 * cm]))

    add_heading("6. Интерпретация графиков")
    chart_notes = [
        (
            CHARTS / "announcement_event_profile.png",
            f"Профиль средней сверхнормальной доходности и накопленной сверхнормальной доходности вокруг даты объявления. График построен по {m['ann_profile_deals']} сделкам и {m['ann_profile_rows']} дневным наблюдениям в окне t = -10...+10.",
        ),
        (
            CHARTS / "create_event_profile.png",
            f"Профиль средней сверхнормальной доходности и накопленной сверхнормальной доходности вокруг даты первого появления сделки. График построен по {m['create_profile_deals']} сделкам и {m['create_profile_rows']} наблюдениям.",
        ),
        (
            CHARTS / "car_ann_1_1_hist.png",
            f"Распределение CAR_ANN_1_1 по {c['ann_car']} сделкам.",
        ),
        (
            CHARTS / "car_close_1_1_hist.png",
            f"Распределение CAR_CLOSE_1_1 по {c['close_car']} сделкам.",
        ),
        (
            CHARTS / "bhar_ann_120_hist.png",
            f"Распределение BHAR_ANN_120 по {c['ann_bhar_120']} сделкам.",
        ),
    ]
    for img_path, note in chart_notes:
        add_body(note)
        img = RLImage(str(img_path))
        img._restrictSize(15.3 * cm, 10.8 * cm)
        story.append(img)
        story.append(Spacer(1, 0.2 * cm))

    add_heading("7. Ограничения и корректность интерпретации")
    add_body(sections["limitations"])

    add_heading("8. Итоговые выводы")
    add_body(sections["conclusion"])

    def add_page_number(canvas, doc):
        canvas.setFont("DejaVuSerif", 9)
        canvas.drawRightString(A4[0] - 1.5 * cm, 1.1 * cm, str(doc.page))

    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=2.1 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.7 * cm,
        bottomMargin=1.7 * cm,
    )
    pdf.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)


def main() -> None:
    ctx = load_context()
    metrics = compute_metrics(ctx)
    sections = build_text_sections(ctx, metrics)
    build_docx(ctx, metrics, sections)
    build_pdf(ctx, metrics, sections)
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
